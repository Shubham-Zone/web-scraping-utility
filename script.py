# Importing the required modules
import shutil
from googlesearch import search
import requests
from bs4 import BeautifulSoup
import csv
import pandas as pd
import os
from docx import Document
# !pip install python-docx


# Function to fetch links from Google search results
def fetch_links(query, num_links=20):
    links = []
    for link in search(query, num=num_links, stop=num_links, pause=2):
        links.append(link)
    return links

# Function to scrape HTML tables and convert them to CSV
def scrape_tables_and_convert_to_csv(link):
    path = requests.get(link)

    soup = BeautifulSoup(path.text, 'html.parser')

    # Check if tables are found
    tables = soup.find_all("table")
    if not tables:
        print("No tables found on", link)
        return

    # Iterate over tables
    for table in tables:
        # Find the header row
        header = table.find("tr")
        if header is not None:
            # Extract header text
            list_header = [item.get_text() for item in header.find_all("th")]

            # Check if the table has multiple columns
            if len(list_header) > 1:
                # Extract data rows
                data_rows = table.find_all("tr")[1:]
                data = []
                for row in data_rows:
                    sub_data = [cell.get_text() for cell in row.find_all(["td", "th"])]
                    data.append(sub_data)

                # Print the header and first few rows of data
                print("Header:", list_header)
                print("First few rows of data:", data[:3])

                # Storing the data into Pandas DataFrame
                dataFrame = pd.DataFrame(data=data, columns=list_header)

                # Converting Pandas DataFrame into CSV file
                dataFrame.to_csv('demo.csv')
                print("CSV file saved successfully.")
                return  # Exit after processing the first table

    print("No suitable tables found for conversion to CSV on", link)

# Function to scrape paragraphs and append them to a Word file
def scrape_paragraphs_and_append_to_word(links, document):
    for link in links:
        try:
            path = requests.get(link)
            soup = BeautifulSoup(path.text, 'html.parser')

            # Get the title of the webpage as the heading
            title = soup.title.string

            # Add heading to document
            document.add_heading(title, level=1)

            # Find all paragraphs
            paragraphs = soup.find_all('p')
            if paragraphs:
                for paragraph in paragraphs:
                    document.add_paragraph(paragraph.get_text())
        except Exception as e:
            print("Error occurred while scraping:", e)

# Main function
def main():
    query = input("Enter your search query: ")
    links = fetch_links(query)
    for link in links:
        print("Scraping tables from:", link)
        scrape_tables_and_convert_to_csv(link)
    document = Document()
    print("Scraping the paragraphs from web ")
    scrape_paragraphs_and_append_to_word(links, document)
    document.save('paragraphs.docx')
    print("Word file saved successfully.")
    # shutil.copyfile("demo.csv", "/content/drive/MyDrive/Colab Notebooks/demo.csv")
    shutil.copyfile("paragraphs.docx", "/content/drive/MyDrive/Colab Notebooks/paragraphs.docx")

if __name__ == "__main__":
    main()
